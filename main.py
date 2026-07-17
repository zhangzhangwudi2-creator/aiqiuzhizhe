"""FastAPI backend for AI Resume Assistant"""
import os
import json
import io
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv
from pypdf import PdfReader
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from prompts import SYSTEM_PROMPT, REWRITE_PROMPT, build_prompt, build_rewrite_prompt

load_dotenv()

app = FastAPI(title="AI 求职助手")

app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

app.mount("/static", StaticFiles(directory="static"), name="static")

client = OpenAI(api_key=os.getenv("DEEPSEEK_API_KEY"), base_url="https://api.deepseek.com")


@app.get("/", response_class=HTMLResponse)
async def index():
    with open("static/index.html", "r", encoding="utf-8") as f:
        content = f.read()
    from fastapi.responses import Response
    return Response(content=content, media_type="text/html", headers={"Cache-Control": "no-cache, no-store, must-revalidate", "Pragma": "no-cache", "Expires": "0"})


def _parse_resume(resume: UploadFile) -> str:
    if not resume.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="请上传 PDF 格式的简历")
    try:
        contents = resume.file.read()
        reader = PdfReader(io.BytesIO(contents))
        text = "".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF 解析失败: {e}")
    if not text.strip():
        raise HTTPException(status_code=400, detail="无法从 PDF 中提取文本")
    return text[:30000]


@app.post("/analyze")
async def analyze(resume: UploadFile = File(...), jd_text: str = Form(...)):
    if not jd_text.strip():
        raise HTTPException(status_code=400, detail="请输入岗位描述")
    resume_text = _parse_resume(resume)
    jd_truncated = jd_text[:15000]
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": build_prompt(resume_text, jd_truncated)}
            ],
            response_format={"type": "json_object"},
            temperature=0.3,
            max_tokens=4096,
        )
        return json.loads(response.choices[0].message.content)
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="AI 响应格式异常")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 调用失败: {e}")


@app.post("/rewrite-resume")
async def rewrite_resume(resume: UploadFile = File(...), jd_text: str = Form(...)):
    if not jd_text.strip():
        raise HTTPException(status_code=400, detail="请输入岗位描述")
    resume_text = _parse_resume(resume)
    jd_truncated = jd_text[:15000]
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": REWRITE_PROMPT},
                {"role": "user", "content": build_rewrite_prompt(resume_text, jd_truncated)}
            ],
            temperature=0.5,
            max_tokens=8192,
        )
        return {"rewritten_resume": response.choices[0].message.content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"优化失败: {e}")


@app.post("/download-rewrite")
async def download_rewrite(resume: UploadFile = File(...), jd_text: str = Form(...)):
    """生成并下载 Word 版优化简历"""
    if not jd_text.strip():
        raise HTTPException(status_code=400, detail="请输入岗位描述")
    resume_text = _parse_resume(resume)
    jd_truncated = jd_text[:15000]

    # 调用 DeepSeek 生成优化简历
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": REWRITE_PROMPT},
                {"role": "user", "content": build_rewrite_prompt(resume_text, jd_truncated)}
            ],
            temperature=0.5,
            max_tokens=8192,
        )
        md_text = response.choices[0].message.content
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"优化失败: {e}")

    # 生成 Word 文档
        # Generate Word document
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3
    style.font.color.rgb = RGBColor(0, 0, 0)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.name = "Arial"

    def clean(text):
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        text = re.sub(r'_([^_]+?)_', r'\1', text)
        text = re.sub(r'`(.+?)`', r'\1', text)
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        text = re.sub(r'~~~+', '', text)
        text = re.sub(r'```', '', text)
        return text

    for line in md_text.split("\n"):
        ls = line.strip()
        if not ls:
            doc.add_paragraph("")
        elif ls.startswith("### "):
            doc.add_heading(clean(ls[4:]), level=3)
        elif ls.startswith("## "):
            doc.add_heading(clean(ls[3:]), level=2)
        elif ls.startswith("# "):
            doc.add_heading(clean(ls[2:]), level=1)
        elif ls.startswith("- ") or ls.startswith("* "):
            doc.add_paragraph(clean(ls[2:]), style="List Bullet")
        elif re.match(r'^\d+\.\s', ls):
            num_content = re.sub(r'^\d+\.\s', '', ls)
            doc.add_paragraph(clean(num_content), style="List Number")
        else:
            doc.add_paragraph(clean(ls))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=optimized_resume.docx"},
    )


@app.get("/health")
async def health():
    return {"status": "ok"}


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8000"))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)
